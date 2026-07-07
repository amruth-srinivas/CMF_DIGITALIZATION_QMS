"""Generate inspection report Word documents with python-docx."""
from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

REPORT_FONT = "Aptos"
_CMF_DIR = Path(__file__).resolve().parent.parent
_REPORT_LOGO_PATH = _CMF_DIR / "cmti (2).png"


def _set_run_font(run, *, size: float, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = REPORT_FONT
    run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), REPORT_FONT)
    r_fonts.set(qn("w:hAnsi"), REPORT_FONT)
    r_fonts.set(qn("w:cs"), REPORT_FONT)
    r_pr.insert(0, r_fonts)


def _set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            element = OxmlElement(tag)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))
            tc_pr.append(element)


def _border_all_sides() -> Dict[str, Dict[str, str]]:
    side = {"val": "single", "sz": "4", "color": "000000", "space": "0"}
    return {"top": side, "left": side, "bottom": side, "right": side}


def _write_cell(
    cell,
    text: str = "",
    *,
    bold: bool = False,
    align: str = "center",
    bg: str | None = None,
    size: float = 10,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text or "")
    _set_run_font(run, size=size, bold=bold)
    if align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if bg:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), bg)
        cell._tc.get_or_add_tcPr().append(shading)
    _set_cell_border(cell, **_border_all_sides())
    # Allow cell to grow vertically with wrapped text
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:noWrap"):
            tc_pr.remove(child)


def _add_logo_to_cell(cell, *, width: float = 1.1) -> None:
    cell.text = ""
    if not _REPORT_LOGO_PATH.is_file():
        return
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    run.add_picture(str(_REPORT_LOGO_PATH), width=Inches(width))
    _set_cell_border(cell, **_border_all_sides())


def _meta_col_spans(total_cols: int) -> tuple[int, int, int]:
    if total_cols == 13:
        return (4, 5, 4)
    span_a = max(1, round((total_cols * 4) / 13))
    span_b = max(1, round((total_cols * 5) / 13))
    span_c = max(1, total_cols - span_a - span_b)
    return span_a, span_b, span_c


def _write_meta_combined_row(
    table,
    row_idx: int,
    fields: tuple[tuple[str, str], ...],
    total_cols: int,
) -> None:
    col = 0
    for (label, value), span in zip(fields, _meta_col_spans(total_cols)):
        end = min(total_cols - 1, col + span - 1)
        if end > col:
            _merge_row_span(table, row_idx, col, end)
        text = f"{label} {value}".strip()
        _write_cell(table.cell(row_idx, col), text, bold=False, align="left", size=11)
        col += span


_REPORT_COL_WIDTHS_13 = [8, 8, 15, 7, 8, 7, 6, 6, 7, 7, 7, 8, 6]
_REPORT_COL_WIDTHS_14 = [5.5, 9.9, 7.7, 7.7, 7.7, 6.6, 6.6, 6.6, 7.7, 7.7, 6.6, 6.6, 6.6, 6.5]


def _report_col_widths(total_cols: int, is_consolidated: bool, max_samples: int) -> List[float]:
    if not is_consolidated and max_samples == 3 and total_cols == 13:
        return list(_REPORT_COL_WIDTHS_13)
    if is_consolidated and max_samples == 3 and total_cols == 14:
        return list(_REPORT_COL_WIDTHS_14)
    base = 100 / total_cols
    return [base] * total_cols


def _split_widths_into_thirds(col_widths: List[float]) -> tuple[int, int, int]:
    total = sum(col_widths)
    t1 = total / 3
    t2 = (2 * total) / 3
    cum = 0.0
    end1 = 0
    for i, w in enumerate(col_widths):
        cum += w
        end1 = i + 1
        if cum >= t1 - 0.001:
            break
    end2 = end1
    for i in range(end1, len(col_widths)):
        cum += col_widths[i]
        end2 = i + 1
        if cum >= t2 - 0.001:
            break
    return end1, end2 - end1, len(col_widths) - end2


def _sign_split_by_width(col_widths: List[float]) -> tuple[int, int]:
    total = sum(col_widths)
    cum = 0.0
    left = 0
    for i, w in enumerate(col_widths):
        cum += w
        left = i + 1
        if cum >= total / 2 - 0.001:
            break
    return left, len(col_widths) - left


def _footer_label_width(span: int) -> int:
    return max(1, min(2, round(span / 2)))


def _write_footer_detail_row(
    table,
    row_idx: int,
    labels: tuple[str, str, str],
    values: tuple[str, str, str],
    chem_cols: int,
    ult_cols: int,
    hard_cols: int,
) -> None:
    label_w = _footer_label_width(chem_cols)
    label_w2 = _footer_label_width(ult_cols)
    label_w3 = _footer_label_width(hard_cols)
    fill1 = max(1, chem_cols - label_w)
    fill2 = max(1, ult_cols - label_w2)
    fill3 = max(1, hard_cols - label_w3)
    segments = [
        (label_w, labels[0], True),
        (fill1, values[0], False),
        (label_w2, labels[1], True),
        (fill2, values[1], False),
        (label_w3, labels[2], True),
        (fill3, values[2], False),
    ]
    col = 0
    for span, text, is_label in segments:
        end = col + span - 1
        if end > col:
            _merge_row_span(table, row_idx, col, end)
        _write_cell(
            table.cell(row_idx, col),
            text,
            bold=is_label,
            align="left",
            bg=None,
            size=10,
        )
        col += span


def _write_footer_sign_row(
    table,
    row_idx: int,
    sign_left: int,
    total_cols: int,
    *,
    inspected_by: str = "",
    checked_by: str = "",
) -> None:
    left_text = f"Inspected by: {inspected_by}".strip() if inspected_by else "Inspected by:"
    right_text = f"Checked by: {checked_by}".strip() if checked_by else "Checked by:"
    _merge_row_span(table, row_idx, 0, sign_left - 1)
    _write_cell(table.cell(row_idx, 0), left_text, bold=True, align="left", size=10)
    _merge_row_span(table, row_idx, sign_left, total_cols - 1)
    _write_cell(table.cell(row_idx, sign_left), right_text, bold=True, align="left", size=10)


def _merge_row_span(table, row_idx: int, col_start: int, col_end: int) -> None:
    if col_end <= col_start:
        return
    a = table.cell(row_idx, col_start)
    b = table.cell(row_idx, col_end)
    a.merge(b)


def _set_table_header_row(table, row_idx: int) -> None:
    """Repeat this row at the top of each printed page when the table spans pages."""
    tr = table.rows[row_idx]._tr
    tr_pr = tr.get_or_add_trPr()
    for child in list(tr_pr):
        if child.tag == qn("w:tblHeader"):
            return
    tr_pr.append(OxmlElement("w:tblHeader"))


def _collect_all_report_rows(payload: Dict[str, Any]) -> List[Dict[str, Any]]:
    rows = list(payload.get("rows") or [])
    if rows:
        return rows
    sheets = payload.get("sheets")
    if sheets:
        collected: List[Dict[str, Any]] = []
        for sheet in sheets:
            collected.extend(sheet.get("rows") or [])
        if collected:
            return collected
    collected = []
    for page in payload.get("pages") or []:
        collected.extend(page.get("rows") or [])
    return collected


def _report_sheet_config(payload: Dict[str, Any], rows: List[Dict[str, Any]]) -> Dict[str, Any]:
    is_consolidated = bool(payload.get("isConsolidated"))
    if is_consolidated:
        max_samples = int(payload.get("quantityCount") or payload.get("maxSamples") or 3)
    else:
        max_samples = max([3] + [len(r.get("measurements") or []) for r in rows], default=3)
    total_cols = int(payload.get("totalCols") or 13)
    return {
        "rows": rows,
        "maxSamples": max_samples,
        "totalCols": total_cols,
        "sheet": "1 of 1",
        "totalQuantity": payload.get("totalQuantity", ""),
        "footerRows": payload.get("footerRows"),
        "inspectedBy": payload.get("inspectedBy"),
        "checkedBy": payload.get("checkedBy"),
        "showFooter": True,
        "isConsolidated": is_consolidated,
        "quantityCount": payload.get("quantityCount"),
    }


def _append_report_banner(doc: Document, payload: Dict[str, Any]) -> None:
    banner = doc.add_table(rows=1, cols=2)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner.autofit = False
    banner.width = Inches(7.1)
    _add_logo_to_cell(banner.cell(0, 0))
    title = str(payload.get("reportTitle") or "INSPECTION REPORT")
    _write_cell(banner.cell(0, 1), title, bold=True, align="center", size=14 if "FINAL" in title.upper() else 16)


def _append_report_sheet_table(doc: Document, payload: Dict[str, Any], sheet: Dict[str, Any]) -> None:
    rows: List[Dict[str, Any]] = sheet.get("rows") or []
    is_consolidated = bool(sheet.get("isConsolidated") or payload.get("isConsolidated"))
    quantity_count = int(
        sheet.get("quantityCount")
        or payload.get("quantityCount")
        or (len(rows[0].get("qtyAverages") or []) if rows else 0)
        or 1
    )
    max_samples = quantity_count if is_consolidated else int(sheet.get("maxSamples") or payload.get("maxSamples") or 3)
    total_cols = int(sheet.get("totalCols") or payload.get("totalCols") or 13)
    show_footer = sheet.get("showFooter", True)
    header_rows = 1 if is_consolidated else 2
    footer_row_count = 6 if show_footer else 0
    table_rows = 3 + header_rows + len(rows) + footer_row_count

    table = doc.add_table(rows=table_rows, cols=total_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = True
    table.width = Inches(7.1)

    r = 0
    meta = [
        (
            ("Report No :", payload.get("reportNo", "")),
            ("Component Title:", payload.get("componentTitle", "")),
            ("Date:", payload.get("date", "")),
        ),
        (
            ("Project No.:", payload.get("projectNo", "")),
            ("Drg No:", payload.get("drgNo", "")),
            ("Sheet", sheet.get("sheet", payload.get("sheet", "1 of 1"))),
        ),
        (
            ("Project Name:", payload.get("projectName", "")),
            ("Quantity:", sheet.get("totalQuantity", payload.get("totalQuantity", ""))),
            ("Assembly", payload.get("assembly", "")),
        ),
    ]
    for fields in meta:
        _write_meta_combined_row(table, r, fields, total_cols)
        r += 1

    col = 0
    if is_consolidated:
        qty_label = "Qty" if quantity_count >= 4 else "Quantity"
        inst_span = 1 if quantity_count >= 4 else 2
        _write_cell(table.cell(r, col), "Sl No", bold=True, bg="F0F0F0")
        col += 1
        table.cell(r, col).merge(table.cell(r, col + 1))
        _write_cell(table.cell(r, col), "Specified Values", bold=True, bg="F0F0F0")
        col += 2
        _write_cell(table.cell(r, col), "Zone", bold=True, bg="F0F0F0")
        col += 1
        for i in range(quantity_count):
            _write_cell(table.cell(r, col + i), f"{qty_label} {i + 1}", bold=True, bg="F0F0F0")
        col += quantity_count
        if inst_span > 1:
            table.cell(r, col).merge(table.cell(r, col + inst_span - 1))
        _write_cell(table.cell(r, col), "Instrument", bold=True, bg="F0F0F0")
        col += inst_span
        rem_span = total_cols - col
        table.cell(r, col).merge(table.cell(r, col + rem_span - 1))
        _write_cell(table.cell(r, col), "Remarks", bold=True, bg="F0F0F0")
        r += 1
    else:
        table.cell(r, col).merge(table.cell(r + 1, col))
        _write_cell(table.cell(r, col), "Sl No", bold=True, bg="F0F0F0")
        col += 1
        table.cell(r, col).merge(table.cell(r + 1, col + 1))
        _write_cell(table.cell(r, col), "Specified Values", bold=True, bg="F0F0F0")
        col += 2
        table.cell(r, col).merge(table.cell(r + 1, col))
        _write_cell(table.cell(r, col), "Zone", bold=True, bg="F0F0F0")
        col += 1
        _merge_row_span(table, r, col, col + max_samples - 1)
        _write_cell(table.cell(r, col), "Measured Values", bold=True, bg="F0F0F0")
        for i in range(max_samples):
            _write_cell(table.cell(r + 1, col + i), str(i + 1), bold=True, bg="F0F0F0")
        col += max_samples
        inst_span = 2
        table.cell(r, col).merge(table.cell(r + 1, col + inst_span - 1))
        _write_cell(table.cell(r, col), "Instrument", bold=True, bg="F0F0F0")
        col += inst_span
        rem_span = total_cols - col
        table.cell(r, col).merge(table.cell(r, col + rem_span - 1))
        _write_cell(table.cell(r, col), "Remarks", bold=True, bg="F0F0F0")
        r += 2

    column_header_start = 3
    column_header_end = r - 1
    for header_row in range(column_header_start, column_header_end + 1):
        _set_table_header_row(table, header_row)

    for item in rows:
        c = 0
        _write_cell(table.cell(r, c), str(item.get("sno", "")), align="center")
        c += 1
        _merge_row_span(table, r, c, c + 1)
        _write_cell(table.cell(r, c), str(item.get("specified", "")), align="left")
        c += 2
        _write_cell(table.cell(r, c), str(item.get("zone", "")), align="center")
        c += 1
        if is_consolidated:
            averages = item.get("qtyAverages") or []
            inst_span = 1 if quantity_count >= 4 else 2
            for i in range(quantity_count):
                val = averages[i] if i < len(averages) else ""
                _write_cell(table.cell(r, c + i), "" if val is None else str(val), align="center")
            c += quantity_count
            if inst_span > 1:
                _merge_row_span(table, r, c, c + 1)
            _write_cell(table.cell(r, c), str(item.get("instrument", "default")), align="center")
            c += inst_span
        else:
            measurements = item.get("measurements") or []
            for i in range(max_samples):
                val = measurements[i] if i < len(measurements) else ""
                _write_cell(table.cell(r, c + i), "" if val is None else str(val), align="center")
            c += max_samples
            _merge_row_span(table, r, c, c + 1)
            _write_cell(table.cell(r, c), str(item.get("instrument", "default")), align="center")
            c += 2
        _merge_row_span(table, r, c, total_cols - 1)
        _write_cell(table.cell(r, c), str(item.get("remarks", "")), align="left")
        r += 1

    col_widths = _report_col_widths(total_cols, is_consolidated, max_samples)
    chem_cols, ult_cols, hard_cols = _split_widths_into_thirds(col_widths)
    sign_left, _sign_right = _sign_split_by_width(col_widths)

    if not show_footer:
        return

    _merge_row_span(table, r, 0, chem_cols - 1)
    _write_cell(table.cell(r, 0), "Chemical Test", bold=True, align="center", bg=None, size=10)
    _merge_row_span(table, r, chem_cols, chem_cols + ult_cols - 1)
    _write_cell(table.cell(r, chem_cols), "Ultrasonic Test", bold=True, align="center", bg=None, size=10)
    _merge_row_span(table, r, chem_cols + ult_cols, total_cols - 1)
    _write_cell(table.cell(r, chem_cols + ult_cols), "Hardness Test", bold=True, align="center", bg=None, size=10)
    r += 1

    test_rows = [
        ("Date", "Date", "Date"),
        ("Report No", "Report No", "W.O.NO"),
        ("Authoriser", "Authoriser", "Hardness Value"),
        ("Status", "Status", "Status"),
    ]
    footer_rows: List[Dict[str, Any]] = sheet.get("footerRows") or payload.get("footerRows") or []
    for i, (chem, ult, hard) in enumerate(test_rows):
        row_vals = footer_rows[i] if i < len(footer_rows) else {}
        _write_footer_detail_row(
            table,
            r,
            (chem, ult, hard),
            (
                str(row_vals.get("chemical") or ""),
                str(row_vals.get("ultrasonic") or ""),
                str(row_vals.get("hardness") or ""),
            ),
            chem_cols,
            ult_cols,
            hard_cols,
        )
        r += 1

    _write_footer_sign_row(
        table,
        r,
        sign_left,
        total_cols,
        inspected_by=str(sheet.get("inspectedBy") or payload.get("inspectedBy") or ""),
        checked_by=str(sheet.get("checkedBy") or payload.get("checkedBy") or ""),
    )


def build_inspection_report_docx(payload: Dict[str, Any]) -> BytesIO:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.47)
    section.right_margin = Inches(0.47)
    section.top_margin = Inches(0.39)
    section.bottom_margin = Inches(0.47)

    normal = doc.styles["Normal"]
    normal.font.name = REPORT_FONT
    normal.font.size = Pt(10)

    rows = _collect_all_report_rows(payload)
    _append_report_banner(doc, payload)
    sheet = _report_sheet_config(payload, rows)
    _append_report_sheet_table(doc, payload, sheet)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
