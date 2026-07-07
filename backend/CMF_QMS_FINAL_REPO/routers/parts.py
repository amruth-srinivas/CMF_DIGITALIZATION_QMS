from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from sqlalchemy.orm import Session
from typing import List
from sqlalchemy.exc import IntegrityError
from sqlalchemy import func, text
from pydantic import BaseModel
import re, tempfile, os

from DB.database import get_db
from DB.models.oms import (
    Part as PartModel, 
    PartType, 
    Order, 
    OrderPartPriority,
    Operation as OperationModel,
    Document as DocumentModel,
    ToolWithPart as ToolWithPartModel,
    OrderPartsRawMaterialLinked as OrderPartsRawMaterialLinkedModel,
    OperationDocument as OperationDocumentModel,
    OutSourcePartStatus as OutSourcePartStatusModel,
    DocumentExtractedData as DocumentExtractedDataModel,
)
from DB.models.configuration import PokayokeCompletedLog
from DB.models.inventory import RawMaterial
from DB.models.access_control import AccessUser
from DB.schemas.oms import Part, PartCreate, PartUpdate

router = APIRouter(
    prefix="/parts",
    tags=["parts"]
)


def _build_part_maps(db: Session):
    """Fetch PartType, RawMaterial, and AccessUser rows once and return id→value maps."""
    type_map = {pt.id: pt.type_name for pt in db.query(PartType).all()}
    rm_map = {rm.id: rm.material_name for rm in db.query(RawMaterial).all()}
    user_map = {u.id: u.user_name for u in db.query(AccessUser).all()}
    return type_map, rm_map, user_map


def _part_to_dict(part: PartModel, type_map: dict, rm_map: dict, user_map: dict) -> dict:
    return {
        "id": part.id,
        "part_name": part.part_name,
        "part_number": part.part_number,
        "type_id": part.type_id,
        "raw_material_id": part.raw_material_id,
        "part_detail": part.part_detail,
        "assembly_id": part.assembly_id,
        "product_id": part.product_id,
        "user_id": part.user_id,
        # "size": part.size,  # New optional size field
        "qty": part.qty,    # New optional quantity field
        # "size": part.size,  # New optional size field
        "qty": part.qty,    # New optional quantity field
        "type_name": type_map.get(part.type_id),
        "raw_material_name": rm_map.get(part.raw_material_id),
        "user_name": user_map.get(part.user_id) if part.user_id else None,
        "created_at": part.created_at,
        "updated_at": part.updated_at,
    }


@router.post("/", response_model=Part, status_code=status.HTTP_201_CREATED)
def create_part(part: PartCreate, db: Session = Depends(get_db)):
    """Create a new part"""
    db_part = db.query(PartModel).filter(PartModel.part_number == part.part_number).first()
    if db_part:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Part with number {part.part_number} already exists"
        )

    db_part = PartModel(**part.model_dump())
    db.add(db_part)
    db.commit()
    db.refresh(db_part)

    # Automatic OrderPartPriority creation disabled
    # if db_part.product_id and db_part.type_id:
    #     part_type = db.query(PartType).filter(PartType.id == db_part.type_id).first()
    #     if part_type and part_type.type_name and part_type.type_name.lower() == "in-house":
    #         orders = db.query(Order).filter(Order.product_id == db_part.product_id).all()
    #         max_priority = db.query(func.max(OrderPartPriority.priority)).scalar() or 0
    #         for index, order in enumerate(orders):
    #             priority_entry = OrderPartPriority(
    #                 order_id=order.id,
    #                 product_id=db_part.product_id,
    #                 part_id=db_part.id,
    #                 priority=max_priority + 1 + index,
    #             )
    #             db.add(priority_entry)
    #         db.commit()

    type_map, rm_map, user_map = _build_part_maps(db)
    return _part_to_dict(db_part, type_map, rm_map, user_map)


@router.get("/", response_model=List[Part])
def get_parts(user_id: int | None = None, db: Session = Depends(get_db)):
    """Get all parts with type, raw material, and user names. Filter by user_id for module-specific views."""
    query = db.query(PartModel).order_by(PartModel.id.asc())
    if user_id is not None:
        query = query.filter(PartModel.user_id == user_id)
    parts = query.all()
    type_map, rm_map, user_map = _build_part_maps(db)
    return [_part_to_dict(p, type_map, rm_map, user_map) for p in parts]


@router.get("/{part_id}", response_model=Part)
def get_part(part_id: int, db: Session = Depends(get_db)):
    """Get a specific part by ID with type, raw_material, and user names."""
    part = db.query(PartModel).filter(PartModel.id == part_id).first()
    if not part:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Part with id {part_id} not found"
        )
    type_map, rm_map, user_map = _build_part_maps(db)
    return _part_to_dict(part, type_map, rm_map, user_map)


@router.get("/part-number/{part_number}", response_model=Part)
def get_part_by_number(part_number: str, db: Session = Depends(get_db)):
    """Get a specific part by its part number string."""
    part = db.query(PartModel).filter(PartModel.part_number == part_number).first()
    if not part:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Part with number {part_number} not found"
        )
    type_map, rm_map, user_map = _build_part_maps(db)
    return _part_to_dict(part, type_map, rm_map, user_map)


@router.get("/product/{product_id}", response_model=List[Part])
def get_parts_by_product(product_id: int, user_id: int | None = None, db: Session = Depends(get_db)):
    """Get all parts for a specific product. Filter by user_id for module-specific views."""
    query = db.query(PartModel).filter(PartModel.product_id == product_id)
    if user_id is not None:
        query = query.filter(PartModel.user_id == user_id)
    parts = query.all()
    type_map, rm_map, user_map = _build_part_maps(db)
    return [_part_to_dict(p, type_map, rm_map, user_map) for p in parts]


@router.get("/assembly/{assembly_id}", response_model=List[Part])
def get_parts_by_assembly(assembly_id: int, user_id: int | None = None, db: Session = Depends(get_db)):
    """Get all parts for a specific assembly. Filter by user_id for module-specific views."""
    query = db.query(PartModel).filter(PartModel.assembly_id == assembly_id)
    if user_id is not None:
        query = query.filter(PartModel.user_id == user_id)
    parts = query.all()
    type_map, rm_map, user_map = _build_part_maps(db)
    return [_part_to_dict(p, type_map, rm_map, user_map) for p in parts]


@router.get("/type/{type_id}", response_model=List[Part])
def get_parts_by_type(type_id: int, user_id: int | None = None, db: Session = Depends(get_db)):
    """Get all parts of a specific type. Filter by user_id for module-specific views."""
    query = db.query(PartModel).filter(PartModel.type_id == type_id)
    if user_id is not None:
        query = query.filter(PartModel.user_id == user_id)
    parts = query.all()
    type_map, rm_map, user_map = _build_part_maps(db)
    return [_part_to_dict(p, type_map, rm_map, user_map) for p in parts]


@router.put("/{part_id}", response_model=Part)
def update_part(part_id: int, part: PartUpdate, db: Session = Depends(get_db)):
    """Update a part"""
    db_part = db.query(PartModel).filter(PartModel.id == part_id).first()
    if not db_part:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Part with id {part_id} not found"
        )

    update_data = part.model_dump(exclude_unset=True)
    for field, value in update_data.items():
        setattr(db_part, field, value)

    db.commit()
    db.refresh(db_part)
    type_map, rm_map, user_map = _build_part_maps(db)
    return _part_to_dict(db_part, type_map, rm_map, user_map)


@router.delete("/{part_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_part(part_id: int, db: Session = Depends(get_db)):
    """Delete a part and all its related data (priorities, pokayoke logs, operations, documents, etc.)."""
    db_part = db.query(PartModel).filter(PartModel.id == part_id).first()
    if not db_part:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Part with id {part_id} not found"
        )

    try:
        # 1. Delete pokayoke logs for this part
        result = db.execute(
            text(
                "SELECT id FROM configuration.pokayoke_completed_logs "
                "WHERE part_id = :pid"
            ),
            {"pid": part_id},
        )
        log_ids = [row[0] for row in result]
        for log_id in log_ids:
            log_obj = (
                db.query(PokayokeCompletedLog)
                .filter(PokayokeCompletedLog.id == log_id)
                .first()
            )
            if log_obj:
                db.delete(log_obj)
        db.flush()

        # Delete from scheduling.part_schedule_status to avoid FK violation
        db.execute(
            text("DELETE FROM scheduling.part_schedule_status WHERE part_id = :pid"),
            {"pid": part_id}
        )

        # 2. Delete order part priorities
        db.query(OrderPartPriority).filter(OrderPartPriority.part_id == part_id).delete(
            synchronize_session=False
        )

        # 3. Delete operations and their documents/tools
        operations = db.query(OperationModel).filter(OperationModel.part_id == part_id).all()
        operation_ids = [op.id for op in operations]
        if operation_ids:
            # Delete from scheduling.planned_schedule_items to avoid FK violation
            db.execute(
                text("DELETE FROM scheduling.planned_schedule_items WHERE operation_id IN :op_ids"),
                {"op_ids": tuple(operation_ids)}
            )

            db.query(OperationDocumentModel).filter(
                OperationDocumentModel.operation_id.in_(operation_ids)
            ).delete(synchronize_session=False)
            
            db.query(ToolWithPartModel).filter(
                ToolWithPartModel.operation_id.in_(operation_ids)
            ).delete(synchronize_session=False)
            
            db.query(OperationModel).filter(OperationModel.id.in_(operation_ids)).delete(
                synchronize_session=False
            )

        # 4. Delete part documents and their extracted data
        # First delete any extracted data associated with this part
        db.query(DocumentExtractedDataModel).filter(
            DocumentExtractedDataModel.part_id == part_id
        ).delete(synchronize_session=False)

        # Now delete the documents associated with this part
        # 4. Delete part documents and their extracted data
        # First delete any extracted data associated with this part
        db.query(DocumentExtractedDataModel).filter(
            DocumentExtractedDataModel.part_id == part_id
        ).delete(synchronize_session=False)

        # Now delete the documents associated with this part
        db.query(DocumentModel).filter(DocumentModel.part_id == part_id).delete(
            synchronize_session=False
        )

        # 5. Delete out source part status records
        db.query(OutSourcePartStatusModel).filter(
            OutSourcePartStatusModel.part_id == part_id
        ).delete(synchronize_session=False)

        # 6. Delete raw material links
        db.query(OrderPartsRawMaterialLinkedModel).filter(
            OrderPartsRawMaterialLinkedModel.part_id == part_id
        ).delete(synchronize_session=False)

        # 7. Delete component_issues records that reference this part
        db.execute(
            text("DELETE FROM maintenance.component_issues WHERE part_id = :pid"),
            {"pid": part_id}
        )

        # 8. Delete tools with part (that are not associated with operations)
        db.query(ToolWithPartModel).filter(ToolWithPartModel.part_id == part_id).delete(
            synchronize_session=False
        )

        # 9. Delete inventory requests for this part
        # First delete return requests that reference these inventory requests
        db.execute(
            text("DELETE FROM inventory.inventory_return_requests WHERE requested_id IN (SELECT id FROM inventory.inventory_requests WHERE part_id = :pid)"),
            {"pid": part_id}
        )
        
        # Then delete the inventory requests
        db.execute(
            text("DELETE FROM inventory.inventory_requests WHERE part_id = :pid"),
            {"pid": part_id}
        )

        # 9. Delete inventory requests for this part
        # First delete return requests that reference these inventory requests
        db.execute(
            text("DELETE FROM inventory.inventory_return_requests WHERE requested_id IN (SELECT id FROM inventory.inventory_requests WHERE part_id = :pid)"),
            {"pid": part_id}
        )
        
        # Then delete the inventory requests
        db.execute(
            text("DELETE FROM inventory.inventory_requests WHERE part_id = :pid"),
            {"pid": part_id}
        )

        # Finally, delete the part itself
        db.delete(db_part)
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error deleting part: {str(e)}"
        )
    return None



# ─────────────────────────────────────────────────────────────────────────────
# ADD THESE IMPORTS at the top of parts.py (merge with existing imports)
# ─────────────────────────────────────────────────────────────────────────────
# from fastapi import UploadFile, File
# import subprocess, tempfile, os, re
# from bs4 import BeautifulSoup
# ─────────────────────────────────────────────────────────────────────────────
# ADD THIS ENDPOINT inside parts.py  (after existing routes, before end of file)
# ─────────────────────────────────────────────────────────────────────────────
@router.post("/parse-doc", status_code=status.HTTP_200_OK)
async def parse_parts_doc(file: UploadFile = File(...)):
    """
    Accept a .docx BOM file and return extracted part rows from ALL pages.
 
    Document structure (CMTI-style BOM):
    ─────────────────────────────────────
    One Word table spans the entire document. For every page the layout is:
 
        [data row]          ← part entry
        [data row]          ← part entry
        …
        [HEADER ROW]        ← "Name of Part | No. of Parts | Size | Part (Assy) No.…"
        [footer 1]          Date / Prepared by …
        [footer 2]          Replaced … Superceded …
        [footer 3]          Central Mfg … | Assembly No. …
        [footer 4]          … | <sheet number> …
        [data row page 2]   ← next page begins
 
    Fixed column indices (merged cells just repeat the text):
        col 3  → Size
        col 5  → Material / Standard
        col 6  → No. of Parts   ← ALWAYS non-empty on real part rows
        col 7  → Name of Part
        col 9  → Part (Assy) No., Drg. Size
    """
    from docx import Document
 
    suffix = os.path.splitext(file.filename or "")[-1].lower()
    if suffix != ".docx":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only .docx files are supported. Please convert .doc to .docx first.",
        )
 
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name
 
    try:
        doc = Document(tmp_path)
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Could not open document: {exc}",
        )
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
 
    if not doc.tables:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="No tables found in the document.",
        )
 
    COL_SIZE     = 3
    COL_MATERIAL = 5
    COL_QTY      = 6   # "No. of Parts"
    COL_NAME     = 7   # "Name of Part"
    COL_PARTNO   = 9   # "Part (Assy) No., Drg. Size"
    FOOTER_ROWS  = 4   # rows immediately after each header that are page footer
 
    def _is_header(cells: list) -> bool:
        joined = " ".join(c.lower() for c in cells)
        return "name of part" in joined and (
            "no. of parts" in joined or "no of parts" in joined
        )
 
    def _cell(row_cells: list, idx: int) -> str:
        if idx >= len(row_cells):
            return ""
        return re.sub(r"\s+", " ", row_cells[idx]).strip()
 
    parts: list[dict] = []
 
    for table in doc.tables:
        all_rows = [[c.text for c in row.cells] for row in table.rows]
 
        header_indices = [i for i, r in enumerate(all_rows) if _is_header(r)]
        if not header_indices:
            continue  # not a BOM table
 
        # Mark header + its 4 footer rows as skip
        skip: set[int] = set()
        for hi in header_indices:
            for offset in range(FOOTER_ROWS + 1):
                skip.add(hi + offset)
 
        for ri, row_cells in enumerate(all_rows):
            if ri in skip:
                continue
 
            part_name   = _cell(row_cells, COL_NAME)
            part_number = _cell(row_cells, COL_PARTNO)
            qty_raw     = _cell(row_cells, COL_QTY)
 
            # 1. Both name and number empty → spacer row
            if not part_name and not part_number:
                continue
 
            # 2. Only one non-empty cell in whole row → row-number spacer
            if len([c for c in row_cells if c.strip()]) == 1:
                continue
 
            # 3. qty is empty → not a real part row (address block, notes, etc.)
            #    Every real part in a CMTI BOM always has a quantity.
            #    Vendor address/notes rows always have an empty qty column —
            #    this single check is sufficient to exclude all of them.
            if not qty_raw.strip():
                continue
 
            # Parse qty: first integer ("57 (55+2)" → 57, "50\n(49+1)" → 50)
            qty: int | None = None
            m = re.search(r"\d+", qty_raw.replace(",", ""))
            if m:
                qty = int(m.group())
 
            parts.append(
                {
                    "part_name":         part_name,
                    "part_number":       part_number,
                    "qty":               qty,
                    # "size":              _cell(row_cells, COL_SIZE) or None,
                    "raw_material_name": _cell(row_cells, COL_MATERIAL) or None,
                    "type_id":           1,    # default: In-house; user can change in UI
                    "part_detail":       None,
                }
            )
 
    if not parts:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=(
                "No part rows could be extracted. "
                "Ensure the BOM table has columns: "
                "'Name of Part', 'No. of Parts', 'Size', 'Part (Assy) No., Drg. Size'."
            ),
        )
 
    return {"parts": parts, "count": len(parts)}
 
 
# ══════════════════════════════════════════════════════════════════════════════
#  2.  NEW  /bulk  — create many parts in a single POST request
# ══════════════════════════════════════════════════════════════════════════════
class BulkPartCreateItem(BaseModel):
    """One part inside a bulk-create request."""
    part_name:       str
    part_number:     str
    type_id:         int        = 1
    raw_material_id: int | None = None
    part_detail:     str | None = None
    assembly_id:     int | None = None
    product_id:      int | None = None
    user_id:         int | None = None
    # size:            str | None = None
    qty:             int | None = None
 
 
class BulkPartCreateRequest(BaseModel):
    parts: list[BulkPartCreateItem]
 
 
class BulkPartCreateResult(BaseModel):
    created:    list[dict]   # successfully created parts
    duplicates: list[str]    # part_numbers that already existed (skipped)
    errors:     list[dict]   # [{"part_number": …, "error": …}]
 
 
@router.post("/bulk", response_model=BulkPartCreateResult, status_code=status.HTTP_200_OK)
def bulk_create_parts(payload: BulkPartCreateRequest, db: Session = Depends(get_db)):
    """
    Create multiple parts in a single database transaction (one HTTP call).
 
    Instead of calling POST /parts/ N times, the frontend sends all rows here
    at once.  Returns three lists so the UI can show per-row status:
      • created    – parts successfully inserted
      • duplicates – part_numbers that already existed (skipped, not a fatal error)
      • errors     – parts that failed for any other reason
    """
    created:    list[dict] = []
    duplicates: list[str]  = []
    errors:     list[dict] = []
 
    type_map, rm_map, user_map = _build_part_maps(db)
 
    # Pre-check all part numbers in ONE query to avoid N round-trips
    incoming_numbers = [item.part_number for item in payload.parts if item.part_number]
    existing_numbers: set[str] = set()
    if incoming_numbers:
        existing_numbers = {
            row.part_number
            for row in db.query(PartModel.part_number)
            .filter(PartModel.part_number.in_(incoming_numbers))
            .all()
        }
 
    for item in payload.parts:
        # Intra-batch duplicate check (catches repeated numbers within the same payload)
        if item.part_number in existing_numbers:
            duplicates.append(item.part_number)
            continue
 
        try:
            db_part = PartModel(**item.model_dump())
            db.add(db_part)
            db.flush()   # assign id without committing yet
            created.append(_part_to_dict(db_part, type_map, rm_map, user_map))
            existing_numbers.add(item.part_number)   # prevent intra-batch dupes
        except Exception as exc:
            db.rollback()
            errors.append({"part_number": item.part_number, "error": str(exc)})
            type_map, rm_map, user_map = _build_part_maps(db)
 
    # Single commit for all successful inserts
    if created:
        try:
            db.commit()
        except Exception as exc:
            db.rollback()
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Commit failed: {exc}",
            )
 
    return BulkPartCreateResult(created=created, duplicates=duplicates, errors=errors)


class BulkDeleteResult(BaseModel):
    assembly_id:   int
    deleted_count: int
    part_ids:      list[int]


@router.delete(
    "/bulk-by-assembly/{assembly_id}",
    response_model=BulkDeleteResult,
    status_code=status.HTTP_200_OK,
)
def bulk_delete_parts_by_assembly(assembly_id: int, db: Session = Depends(get_db)):
    """
    Delete ALL parts linked to the given assembly_id in a single call,
    along with every dependent record for each part — identical cleanup
    logic to the single-part DELETE /{part_id} endpoint.

    Called from the Assembly Document Panel when the user clicks
    "Delete All Parts" for an assembly.

    Returns the count and IDs of deleted parts so the frontend can
    update its local state without a refetch.
    """
    # ── 1. Find all part IDs for this assembly ────────────────────────────────
    parts = (
        db.query(PartModel)
        .filter(PartModel.assembly_id == assembly_id)
        .all()
    )

    if not parts:
        # Nothing to delete — return gracefully (not a 404)
        return BulkDeleteResult(
            assembly_id=assembly_id,
            deleted_count=0,
            part_ids=[],
        )

    part_ids = [p.id for p in parts]

    try:
        # ── 2. Pokayoke logs ──────────────────────────────────────────────────
        for part_id in part_ids:
            result = db.execute(
                text(
                    "SELECT id FROM configuration.pokayoke_completed_logs "
                    "WHERE part_id = :pid"
                ),
                {"pid": part_id},
            )
            log_ids = [row[0] for row in result]
            for log_id in log_ids:
                log_obj = (
                    db.query(PokayokeCompletedLog)
                    .filter(PokayokeCompletedLog.id == log_id)
                    .first()
                )
                if log_obj:
                    db.delete(log_obj)
        db.flush()

        # ── 3. Scheduling: part_schedule_status ──────────────────────────────
        db.execute(
            text(
                "DELETE FROM scheduling.part_schedule_status "
                "WHERE part_id = ANY(:pids)"
            ),
            {"pids": part_ids},
        )

        # ── 4. Order part priorities ──────────────────────────────────────────
        db.query(OrderPartPriority).filter(
            OrderPartPriority.part_id.in_(part_ids)
        ).delete(synchronize_session=False)

        # ── 5. Operations → planned_schedule_items, documents, tools ─────────
        operations = (
            db.query(OperationModel)
            .filter(OperationModel.part_id.in_(part_ids))
            .all()
        )
        operation_ids = [op.id for op in operations]

        if operation_ids:
            db.execute(
                text(
                    "DELETE FROM scheduling.planned_schedule_items "
                    "WHERE operation_id = ANY(:oids)"
                ),
                {"oids": operation_ids},
            )
            db.query(OperationDocumentModel).filter(
                OperationDocumentModel.operation_id.in_(operation_ids)
            ).delete(synchronize_session=False)
            db.query(ToolWithPartModel).filter(
                ToolWithPartModel.operation_id.in_(operation_ids)
            ).delete(synchronize_session=False)
            db.query(OperationModel).filter(
                OperationModel.id.in_(operation_ids)
            ).delete(synchronize_session=False)

        # ── 6. Part documents + extracted data ───────────────────────────────
        db.query(DocumentExtractedDataModel).filter(
            DocumentExtractedDataModel.part_id.in_(part_ids)
        ).delete(synchronize_session=False)
        db.query(DocumentModel).filter(
            DocumentModel.part_id.in_(part_ids)
        ).delete(synchronize_session=False)

        # ── 7. OutSource part status ──────────────────────────────────────────
        db.query(OutSourcePartStatusModel).filter(
            OutSourcePartStatusModel.part_id.in_(part_ids)
        ).delete(synchronize_session=False)

        # ── 8. Raw material links ─────────────────────────────────────────────
        db.query(OrderPartsRawMaterialLinkedModel).filter(
            OrderPartsRawMaterialLinkedModel.part_id.in_(part_ids)
        ).delete(synchronize_session=False)

        # ── 9. Maintenance component issues ──────────────────────────────────
        db.execute(
            text(
                "DELETE FROM maintenance.component_issues "
                "WHERE part_id = ANY(:pids)"
            ),
            {"pids": part_ids},
        )

        # ── 10. Tools with part (not operation-linked) ────────────────────────
        db.query(ToolWithPartModel).filter(
            ToolWithPartModel.part_id.in_(part_ids)
        ).delete(synchronize_session=False)

        # ── 11. Inventory return requests + inventory requests ────────────────
        db.execute(
            text(
                "DELETE FROM inventory.inventory_return_requests "
                "WHERE requested_id IN ("
                "  SELECT id FROM inventory.inventory_requests "
                "  WHERE part_id = ANY(:pids)"
                ")"
            ),
            {"pids": part_ids},
        )
        db.execute(
            text(
                "DELETE FROM inventory.inventory_requests "
                "WHERE part_id = ANY(:pids)"
            ),
            {"pids": part_ids},
        )

        # ── 12. Finally delete all parts ──────────────────────────────────────
        db.query(PartModel).filter(
            PartModel.id.in_(part_ids)
        ).delete(synchronize_session=False)

        db.commit()

    except Exception as exc:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Bulk delete failed: {str(exc)}",
        )

    return BulkDeleteResult(
        assembly_id=assembly_id,
        deleted_count=len(part_ids),
        part_ids=part_ids,
    )